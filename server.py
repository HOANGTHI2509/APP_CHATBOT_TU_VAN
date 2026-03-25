import os
import io
import base64
import requests
import shutil
import time
import uuid
import re
import sqlite3
import hashlib
import chromadb
import fitz  # PyMuPDF
from PIL import Image, ImageEnhance
from fastapi import FastAPI, File, UploadFile, Form, Request, Depends, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from dotenv import load_dotenv, find_dotenv
from pydantic import BaseModel
from langchain_chroma import Chroma
from langchain_openai import OpenAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter

app = FastAPI()

# --- KHỞI TẠO DATABASE SQLITE ---
DB_FILE = "dnu_chatbot.db"

def init_db():
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    # Bảng Tài khoản
    c.execute('''CREATE TABLE IF NOT EXISTS users (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    username TEXT UNIQUE,
                    password_hash TEXT,
                    role TEXT DEFAULT 'student'
                 )''')
    # Bảng Phiên Chat (Cuộc trò chuyện)
    c.execute('''CREATE TABLE IF NOT EXISTS chats (
                    id TEXT PRIMARY KEY,
                    user_id INTEGER,
                    title TEXT,
                    created_at DATETIME DEFAULT CURRENT_TIMESTAMP
                 )''')
    # Bảng Tin nhắn
    c.execute('''CREATE TABLE IF NOT EXISTS messages (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    chat_id TEXT,
                    sender TEXT,
                    text TEXT,
                    timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
                 )''')
    
    # Tạo tài khoản Admin mặc định nếu chưa có
    c.execute("SELECT * FROM users WHERE username='admin'")
    if not c.fetchone():
        admin_pass = hashlib.sha256("admin123".encode()).hexdigest()
        c.execute("INSERT INTO users (username, password_hash, role) VALUES (?, ?, ?)", ("admin", admin_pass, "admin"))
    
    conn.commit()
    conn.close()

init_db()

def get_db():
    conn = sqlite3.connect(DB_FILE)
    conn.row_factory = sqlite3.Row
    return conn

# --- MODELS ---
class LoginRequest(BaseModel):
    username: str
    password: str

class RegisterRequest(BaseModel):
    username: str
    password: str

class ScrapeRequest(BaseModel):
    url: str

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

PDF_DIR = "c:/Users/Admin/Downloads/BOTCHAT/PDF"
os.makedirs(PDF_DIR, exist_ok=True)

LANGFLOW_URL = "http://localhost:7861/api/v1/run/c8bcdea3-ccff-4d9c-8bd7-2da19c26e05f"
load_dotenv(find_dotenv())
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
API_KEY_LANGFLOW = os.getenv("LANGFLOW_API_KEY")

def is_text_garbled(text):
    """Kiểm tra độ lỗi font (Chống lại các PDF rác chữ như: IMI HOC DAI NAM, Ha N3i)"""
    if not text or len(text.strip()) < 50: return True
    
    if re.search(r'[ñçüœ£ßåäöæø§©®]', text, re.IGNORECASE): return True
    if re.search(r'[a-zA-Z][&@$^*~][a-zA-Z]', text): return True

    words = text.split()
    if not words: return True
    
    error_words = 0
    for w in words:
        w_clean = w.strip('.,()[]{}":;!?\'-')
        if not w_clean: continue
        if any(c.isdigit() for c in w_clean) and any(c.isalpha() for c in w_clean):
            if not re.match(r'^(\d+[\.,]?\d*[đdkKmMgG\%]?|\d{1,2}/\d{1,2}/\d{2,4}|\d+[a-zA-Z]{1,2})$', w_clean):
                error_words += 1
                continue
        if re.search(r'[a-z][A-Z]', w_clean):
            error_words += 1
            
    # Ngưỡng phát hiện: Chỉ cần 4% chữ là RÁC -> KÍCH HOẠT TRÍ TUỆ NHÂN TẠO VISION
    if error_words / len(words) > 0.04:
        return True
    return False

def ocr_page_with_ai(page):
    """Quét ảnh PDF bằng GPT-4o-mini (Cho ra Markdown bảng biểu hoàn hảo)"""
    try:
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
        img_bytes = pix.tobytes("png")
        img = Image.open(io.BytesIO(img_bytes)).convert('L')
        img = ImageEnhance.Contrast(img).enhance(1.8)
        
        buffer = io.BytesIO()
        img.save(buffer, format="JPEG", quality=85)
        b64 = base64.b64encode(buffer.getvalue()).decode('utf-8')
        
        headers = {"Content-Type": "application/json", "Authorization": f"Bearer {OPENAI_API_KEY}"}
        payload = {
            "model": "gpt-4o-mini",
            "messages": [{"role": "user", "content": [
                {"type": "text", "text": "Đọc nội dung và trả về Markdown. QUY TẮC: Giữ nguyên bảng, nếu 1 ô có nhiều dòng dùng thẻ <br>. KHÔNG lược bỏ text."},
                {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{b64}"}}
            ]}],
            "max_tokens": 2000
        }
        res = requests.post("https://api.openai.com/v1/chat/completions", headers=headers, json=payload, timeout=30)
        return res.json()['choices'][0]['message']['content'].replace("```markdown", "").replace("```", "").strip()
    except Exception as e:
        return f"\n[Lỗi kết nối OpenAI Vision: {str(e)}]\n"

def ocr_page_with_tesseract(page):
    """Quét ảnh PDF bằng công nghệ Tesseract OCR (Offline + Free)"""
    try:
        import pytesseract
        # Lưu ý: Cần cài đặt Tesseract Windows và chỉnh đường dẫn này nếu khác
        pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
        
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
        img_bytes = pix.tobytes("png")
        img = Image.open(io.BytesIO(img_bytes)).convert('L')
        # Tùy chọn psm 6 giúp đọc ảnh chứa nhiều dải văn bản
        text = pytesseract.image_to_string(img, lang="vie", config='--psm 6')
        return text.strip()
    except Exception as e:
        return f"\n[Lỗi Tesseract: Chưa cài đặt Tesseract-OCR trên máy tính hoặc chạy lỗi. Chi tiết: {str(e)}]\n"


@app.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    try:
        file_path = os.path.join(PDF_DIR, file.filename)
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
            
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()

        print(f"Bắt đầu nạp file {file.filename} lên Máy Chủ AI Docker...")
        
        # TRỞ VỀ CHẾ ĐỘ LOCAL AN TOÀN TUYỆT ĐỐI
        embeddings = OpenAIEmbeddings(openai_api_key=OPENAI_API_KEY, model="text-embedding-3-small")
        vectorstore = Chroma(
            persist_directory="./chroma_db",
            collection_name="BOTCHAT", 
            embedding_function=embeddings
        )
        splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        chunks = splitter.split_text(content)
        metadatas = [{"source": file.filename}] * len(chunks)
        vectorstore.add_texts(texts=chunks, metadatas=metadatas)
        print(f"Đã nạp {len(chunks)} chunks vào ChromaDB!")
        
        # Báo cho Langflow bằng 1 câu chat tượng trưng
        payload = {
            "output_type": "chat",
            "input_type": "chat",
            "input_value": f"Thông báo: Đã nạp thành công {file.filename} vào database.",
            "session_id": str(uuid.uuid4())
        }
        headers = {"x-api-key": API_KEY_LANGFLOW}
        response = requests.post(LANGFLOW_URL, json=payload, headers=headers)
        response.raise_for_status()
        
        return {"status": "success", "message": f"Dữ liệu sạch {file.filename} đã nạp!"}
    
    except Exception as e:
        return {"status": "error", "message": str(e)}

@app.post("/scrape")
async def scrape_url(req: ScrapeRequest):
    try:
        res = requests.get(req.url, headers={'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'}, timeout=15)
        res.raise_for_status()
        html_content = res.text
        
        # Dùng BeautifulSoup để bóc tách thông minh
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(res.text, 'html.parser')
        
        # Bước 1: Xóa các thẻ rác cơ bản (thêm head để không dính thẻ title, meta)
        for element in soup(['head', 'script', 'style', 'nav', 'header', 'footer', 'aside', 'noscript', 'button', 'iframe', 'form']):
            element.decompose()
            
        # Bước 2: Thử tìm container chứa bài viết chính
        main_container = soup.find('article') or soup.find('main')
        if not main_container:
            main_container = soup.find('div', class_=re.compile(r'chi-tiet|detail|post-content|article-content|entry-content', re.I))
            
        content_soup = main_container if main_container else soup
        
        # Thử tìm tiêu đề bài viết từ thẻ <title> làm mốc xén Menu
        title_match = re.search(r'<title>(.*?)</title>', res.text, re.IGNORECASE)
        page_title = ""
        if title_match:
            parts = re.split(r'\|', title_match.group(1))
            page_title = parts[-1].strip()

        # Bước 3: Xử lý các bảng (Tables) biến chúng thành danh sách siêu tối ưu cho RAG AI
        for table in content_soup.find_all('table'):
            table_text_list = []
            rows = table.find_all('tr')
            
            headers = []
            current_group = ""
            
            for i, row in enumerate(rows):
                cols = [col.get_text(strip=True) for col in row.find_all(['td', 'th'])]
                if not cols: continue
                
                if i == 0 and len(cols) > 1:
                    headers = cols
                    continue
                
                if len(cols) == 1:
                    current_group = cols[0]
                    table_text_list.append(f"\n=== MỤC/NHÓM: {current_group.upper()} ===")
                else:
                    row_data = []
                    for j, c in enumerate(cols):
                        if not c or c.lower() == 'link' or (headers and j < len(headers) and headers[j].lower() == 'tt'):
                            continue
                        head_name = headers[j] if j < len(headers) else f"Thông tin {j+1}"
                        row_data.append(f"{head_name}: {c}")
                        
                    if row_data:
                        prefix = f"[{current_group}] " if current_group else ""
                        table_text_list.append(f" ✔️ {prefix}" + " | ".join(row_data))
            
            table.replace_with("\n\n" + "\n".join(table_text_list) + "\n\n")

        # Bước 4: Lấy text ra và làm sạch (chống nát format do inline tags)
        for tag in content_soup.find_all(['p', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li', 'article', 'section']):
            if tag.string is None: # không làm hỏng dữ liệu gốc
                pass
            tag.insert_before('\n')
            tag.insert_after('\n')
        for tag in content_soup.find_all('br'):
            tag.replace_with('\n')

        # Dùng khoảng trắng cho inline tags thay vì \n\n, giúp các chữ <b/> không bị rớt dòng
        text = content_soup.get_text(separator=' ')
        
        # Dọn dẹp thảm họa khoảng trắng thừa
        text = re.sub(r'[ \t]+', ' ', text)
        text = re.sub(r' \n|\n ', '\n', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        lines = []
        raw_lines = text.split('\n')
        
        # Cắt đầu: Lấy Tiêu đề từ <title> để chặn chặn Menu
        start_idx = 0
        if page_title:
            for idx, l in enumerate(raw_lines):
                if page_title.lower() in l.lower() or l.lower() in page_title.lower():
                    if len(l.strip()) > 5:
                        start_idx = idx
                        break
                        
        for idx in range(start_idx, len(raw_lines)):
            line = raw_lines[idx].strip()
            if not line: continue
            
            # Cắt đuôi tàn dư: Footer / Related posts
            lower_line = line.lower()
            if any(lower_line.startswith(k) for k in ["bài viết mới nhất", "bài viết liên quan", "sự kiện gần", "đăng ký tư vấn", "xem tất cả"]):
                break
                
            lines.append(line)
        
        # Nối lại nhưng giữ khoảng cách cho block
        final_text = ""
        for line in lines:
            if "✔️" in line or "===" in line:
                final_text += line + "\n"
            else:
                final_text += line + "\n\n"
        
        # Rút trích tiêu đề và tạo tên file duy nhất
        title_match = re.search(r'<title>(.*?)</title>', res.text, re.IGNORECASE)
        base_title = title_match.group(1).strip() if title_match else "Website_Data"
        
        # Lấy slug từ URL để phân biệt
        url_slug = req.url.split('/')[-1] or req.url.split('/')[-2] or "page"
        url_slug = re.sub(r'[\\/*?:"<>|]', "", url_slug)[:20]
        
        timestamp = time.strftime("%H%M%S")
        safe_title = f"{url_slug}_{timestamp}"
        
        return {"status": "success", "title": safe_title, "text": final_text}
    except Exception as e:
        return {"status": "error", "message": f"Website chặn Scraper hoặc Link sai: {str(e)}"}

class IngestRequest(BaseModel):
    text: str
    filename: str

@app.post("/api/ingest")
async def api_ingest(req: IngestRequest):
    try:
        print(f"API Ingest: Bắt đầu nạp {req.filename} vào database cục bộ...")
        # TRỞ VỀ CHẾ ĐỘ LOCAL AN TOÀN TUYỆT ĐỐI
        embeddings = OpenAIEmbeddings(openai_api_key=OPENAI_API_KEY, model="text-embedding-3-small")
        vectorstore = Chroma(
            persist_directory="./chroma_db",
            collection_name="BOTCHAT", 
            embedding_function=embeddings
        )
        splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        chunks = splitter.split_text(req.text)
        metadatas = [{"source": req.filename}] * len(chunks)
        vectorstore.add_texts(texts=chunks, metadatas=metadatas)
        print(f"Đã nạp {len(chunks)} chunks từ file {req.filename}!")
        return {"message": "Success"}
    except Exception as e:
        return {"error": str(e)}

# --- API QUẢN LÝ TÀI KHOẢN ---
@app.post("/login")
async def login(req: LoginRequest):
    conn = get_db()
    user = conn.execute("SELECT * FROM users WHERE username=?", (req.username,)).fetchone()
    conn.close()
    
    if not user:
        return {"status": "error", "message": "Tài khoản không tồn tại!"}
    
    pass_hash = hashlib.sha256(req.password.encode()).hexdigest()
    if user["password_hash"] != pass_hash:
        return {"status": "error", "message": "Sai mật khẩu!"}
        
    return {
        "status": "success", 
        "user": {"id": user["id"], "username": user["username"], "role": user["role"]}
    }

@app.post("/register")
async def register(req: RegisterRequest):
    conn = get_db()
    pass_hash = hashlib.sha256(req.password.encode()).hexdigest()
    try:
        conn.execute("INSERT INTO users (username, password_hash) VALUES (?, ?)", (req.username, pass_hash))
        conn.commit()
        new_id = conn.execute("SELECT last_insert_rowid()").fetchone()[0]
        conn.close()
        return {"status": "success", "user": {"id": new_id, "username": req.username, "role": "student"}}
    except sqlite3.IntegrityError:
        conn.close()
        return {"status": "error", "message": "Tên đăng nhập đã tồn tại!"}

@app.get("/users")
async def get_all_users():
    conn = get_db()
    users = conn.execute("SELECT id, username, role FROM users ORDER BY id DESC").fetchall()
    conn.close()
    return {"status": "success", "users": [dict(u) for u in users]}

@app.delete("/users/{user_id}")
async def delete_user(user_id: int):
    conn = get_db()
    conn.execute("DELETE FROM users WHERE id=?", (user_id,))
    conn.commit()
    conn.close()
    return {"status": "success", "message": "Đã xóa tài khoản"}

@app.put("/users/{user_id}/role")
async def update_user_role(user_id: int, req: Request):
    data = await req.json()
    new_role = data.get("role", "student")
    conn = get_db()
    conn.execute("UPDATE users SET role=? WHERE id=?", (new_role, user_id))
    conn.commit()
    conn.close()
    return {"status": "success", "message": "Đã cập nhật phân quyền"}

# --- API LỊCH SỬ CHAT ---
@app.get("/history/{user_id}")
async def get_history(user_id: int):
    conn = get_db()
    chats = conn.execute("SELECT * FROM chats WHERE user_id=? ORDER BY created_at DESC", (user_id,)).fetchall()
    conn.close()
    return {"status": "success", "chats": [dict(c) for c in chats]}

@app.get("/messages/{chat_id}")
async def get_messages(chat_id: str):
    conn = get_db()
    msgs = conn.execute("SELECT sender, text FROM messages WHERE chat_id=? ORDER BY timestamp ASC", (chat_id,)).fetchall()
    conn.close()
    return {"status": "success", "messages": [dict(m) for m in msgs]}

@app.get("/files")
async def list_files():
    from datetime import datetime
    try:
        if not os.path.exists(PDF_DIR):
            return {"status": "success", "files": []}
        
        file_list = []
        for f in os.listdir(PDF_DIR):
            if f.endswith(".txt"):
                path = os.path.join(PDF_DIR, f)
                stat = os.stat(path)
                mtime = stat.st_mtime
                date_str = datetime.fromtimestamp(mtime).strftime('%Y-%m-%d')
                size_mb = stat.st_size / (1024 * 1024)
                
                file_list.append({
                    "name": f,
                    "date": date_str,
                    "timestamp": mtime,
                    "size": f"{size_mb:.3f} MB"
                })
        
        file_list.sort(key=lambda x: x["timestamp"], reverse=True)
        return {"status": "success", "files": file_list}
        
        file_list.sort(key=lambda x: x["timestamp"], reverse=True)
        return {"status": "success", "files": file_list}
    except Exception as e:
        return {"status": "error", "message": str(e)}

@app.get("/view-file/{filename}")
async def view_file_content(filename: str):
    try:
        path = os.path.join(PDF_DIR, filename)
        if not os.path.exists(path):
            return {"status": "error", "message": "Không tìm thấy file!"}
        with open(path, "r", encoding="utf-8") as f:
            content = f.read()
        return {"status": "success", "content": content}
    except Exception as e:
        return {"status": "error", "message": str(e)}

@app.post("/chat")
async def chat_langflow(req: Request):
    try:
        data_in = await req.json()
        user_message = data_in.get("message", "Xin chào")
        session_id = data_in.get("session_id", str(uuid.uuid4()))
        user_id = data_in.get("user_id", 0) # 0 nếu chưa đăng nhập (khách)
        chat_title = data_in.get("title", user_message[:30] + "...")

        # Lưu vào SQLite
        conn = get_db()
        # Tạo session chat báo nếu chưa có
        existing_chat = conn.execute("SELECT id FROM chats WHERE id=?", (session_id,)).fetchone()
        if not existing_chat and user_id != 0:
            conn.execute("INSERT INTO chats (id, user_id, title) VALUES (?, ?, ?)", (session_id, user_id, chat_title))
        
        # Lưu câu hỏi của User
        conn.execute("INSERT INTO messages (chat_id, sender, text) VALUES (?, ?, ?)", (session_id, "user", user_message))
        conn.commit()

        payload = {
            "output_type": "chat",
            "input_type": "chat",
            "input_value": user_message,
            "session_id": session_id
        }
        headers = {"x-api-key": API_KEY_LANGFLOW}
        res = requests.post(LANGFLOW_URL, json=payload, headers=headers)
        res.raise_for_status()
        
        data_out = res.json()
        try:
            answer = data_out["outputs"][0]["outputs"][0]["results"]["message"]["text"]
        except Exception:
            try:
                answer = data_out["outputs"][0]["outputs"][0]["artifacts"]["message"]
            except:
                answer = str(data_out)
                
        # Lưu câu trả lời của Bot
        conn.execute("INSERT INTO messages (chat_id, sender, text) VALUES (?, ?, ?)", (session_id, "bot", answer))
        conn.commit()
        conn.close()
                
        return {"status": "success", "answer": answer, "session_id": session_id}
    except Exception as e:
        return {"status": "error", "message": f"Python -> Langflow Lỗi: {str(e)}"}

@app.post("/extract")
async def extract_text(
    file: UploadFile = File(...),
    use_ai: str = Form("false"),
    use_tesseract: str = Form("false"),
    auto_ai: str = Form("false")
):
    try:
        content = await file.read()
        if file.filename.endswith(".txt"):
            md_text = content.decode("utf-8", errors="ignore")
        elif file.filename.endswith(".pdf"):
            doc = fitz.open(stream=content, filetype="pdf")
            full_text = []
            for i, page in enumerate(doc):
                # Nếu User chọn Dùng AI (use_ai == true) -> BẮT BUỘC dùng Vision cho tất cả
                if use_ai == "true":
                    ext_text = ocr_page_with_ai(page)
                elif use_tesseract == "true":
                    ext_text = ocr_page_with_tesseract(page)
                else:
                    # Thử trích xuất cơ bản
                    ext_text = page.get_text("text")
                    # Chỉ dùng AI cứu hộ nếu User cho phép (auto_ai == true) và text bị rác
                    if auto_ai == "true" and is_text_garbled(ext_text):
                        ext_text = ocr_page_with_ai(page)
                
                full_text.append(ext_text)
            
            md_text = "\n\n".join(full_text)
            
            # LỌC SẠCH KÝ TỰ TÀNG HÌNH (Control Characters: \x00, \x0c...) GÂY LỖI 400 OPENAI JSON
            # Giữ lại \n, \r, \t, và các ký tự hiển thị được.
            import re
            md_text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', md_text)
            
        elif file.filename.endswith((".doc", ".docx")):
            try:
                import docx
                import io
                doc_word = docx.Document(io.BytesIO(content))
                
                full_text = []
                for para in doc_word.paragraphs:
                    if para.text.strip():
                        full_text.append(para.text)
                
                # Đọc thô cả bảng biểu trong Word
                for table in doc_word.tables:
                    for row in table.rows:
                        row_data = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
                        full_text.append(" | ".join(row_data))
                
                md_text = "\n\n".join(full_text)
            except Exception as e:
                return {"status": "error", "message": f"Định dạng hoặc cấu trúc file bị hỏng! (Lỗi phổ biến với file .doc 97-2003). Vui lòng Save As sang .docx hoặc .pdf.\n\nChi tiết kỹ thuật: {str(e)}"}

        else:
            md_text = "Định dạng không hỗ trợ trích xuất text."
            
        return {"status": "success", "text": md_text}
        
    except Exception as e:
        import traceback
        return {"status": "error", "message": f"Lỗi: {str(e)} \n {traceback.format_exc()}"}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=8000)
